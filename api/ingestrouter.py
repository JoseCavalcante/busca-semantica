from fastapi import APIRouter, File, UploadFile, Form, HTTPException, Depends, BackgroundTasks
import json
import uuid
import traceback
from typing import Optional
from db.models import User, CandidateDocument, CandidateDocument
from api.authenticationrouter import get_current_user
from sqlalchemy.orm import Session
from db.database import get_db

from service.extracaoPDFservice import extract_pdf_metadata, split_text_into_chunks, extract_docx_metadata
from service.enrichmentservice import enrich_resume_data
from service.embeddingservice import embeddingService
from service.upsertservice import upsert_rich_vectors

router = APIRouter()

async def process_file_background(
    file_filename: str, 
    file_content_type: str, 
    file_bytes: bytes, 
    manual_metadata: dict, 
    user_tenant_id: int, 
    document_id: int, 
    db: Session
):
    """
    Background Task: Handles heavy lifting (Extraction, Enrichment, Embedding, Upsert).
    Updates DB status upon completion or error.
    """
    print(f"--- [BG] Starting Processing for Doc ID: {document_id} ---")
    
    # Re-create a temporary file-like object for logic that needs it (if any)
    # Ideally services should accept bytes, but for now we might need to mock UploadFile behaviour or save to temp
    # simpler approach: Modify services to accept bytes if possible, or save to temp.
    # checking extract_pdf_metadata... it uses fitz.open(stream=file.file.read(), filetype="pdf")
    # So we can pass bytes directly if we adapt the service call or mock the file object.
    # Let's try to pass bytes to a modified service wrapper or handle it here.
    
    try:
        # Update Status to PROCESSING
        doc = db.query(CandidateDocument).filter(CandidateDocument.id == document_id).first()
        if doc:
            doc.processing_status = "PROCESSING"
            db.commit()

        # 1. Extract Text
        full_text = ""
        if file_filename.lower().endswith(".pdf"):
            import fitz
            with fitz.open(stream=file_bytes, filetype="pdf") as doc_pdf:
                for page in doc_pdf:
                    full_text += page.get_text()
        elif file_filename.lower().endswith(".docx"):
            import io
            # simple docx extraction from bytes
            # We need to check if extract_docx_metadata can accept bytes. 
            # It expects UploadFile usually. Let's do a localized extraction here to be safe and robust.
            pass # TODO: Implement robust byte handling if needed, or assume PDF for now given the context
            
            # Fallback for DOCX logic reuse or simple implementation:
            # For brevity/robustness in BG task without complex rewrites:
            # If docx support is strictly needed in BG, we need to adapt extract_docx_metadata.
            # Assuming PDF primary for now as per previous checks.
            # If it IS docx, let's try to load it.
            # BUT: extract_docx_metadata was imported. Let's see if we can use it with a SpooledTemporaryFile mock?
            # It's cleaner to just implement the logic here for the 'text' part or refactor service.
            # Let's stick to PDF optimization first or basic text extract.
            
            # HACK: If it's docx, we skip for this iteration of "Background Tasks" implementation unless critical.
            # User uses PDFs mostly in examples. I will handle PDF robustly.
            pass

        if not full_text and (file_filename.lower().endswith(".docx") or file_content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
             # On-the-fly docx
             import docx
             import io
             doc_stream = io.BytesIO(file_bytes)
             doc_obj = docx.Document(doc_stream)
             full_text = "\n".join([p.text for p in doc_obj.paragraphs])

        if not full_text.strip():
             # Try falling back to legacy service if needed or fail
             raise Exception("Could not extract text (empty or unsupported format in BG handler)")

        # Update DB with text early (so keyword search works even if embedding fails later?)
        # doc.full_text = full_text # It was already saved as empty or we update it now?
        # Actually, in the endpoint we saved it? No, in endpoint we don't have text yet!
        # So we MUST update it here.
        doc = db.query(CandidateDocument).filter(CandidateDocument.id == document_id).first()
        doc.full_text = full_text
        db.commit()

        # 2. Enrichment
        enriched_data = enrich_resume_data(full_text)
        
        # 3. Metadata Prep
        base_metadata = {
            "source_file": file_filename,
            "upload_date": manual_metadata.get("upload_date", ""),
            "candidate_id": manual_metadata.get("id_candidato", ""),
            "candidate_name": enriched_data.get("candidato", {}).get("nome_completo", ""),
            "email": (enriched_data.get("candidato", {}).get("email") or [""])[0],
            "processing_status": "COMPLETED", # Logic tag
            "tenant_id": user_tenant_id # For Pinecone
        }
        
        # Update DB Metadata fields
        doc.candidate_name = base_metadata["candidate_name"]
        doc.email = base_metadata["email"]
        db.commit() # Save progress

        # 4. Chunking & Embedding
        chunks = split_text_into_chunks(full_text, chunk_size=500, overlap=50)
        vectors_to_upsert = []
        
        for i, chunk in enumerate(chunks):
            vector_values = await embeddingService(chunk)
            if not vector_values: continue
            
            chunk_id = str(uuid.uuid4())
            chunk_metadata = base_metadata.copy()
            chunk_metadata["chunk_text"] = chunk
            chunk_metadata["chunk_index"] = i
            
            vectors_to_upsert.append({
                "id": chunk_id,
                "values": vector_values,
                "metadata": chunk_metadata
            })
            
        # 5. Upsert
        upsert_rich_vectors(vectors_to_upsert)
        
        # Final Success State
        doc.processing_status = "COMPLETED"
        db.commit()
        print(f"--- [BG] Processing Finished Success: {document_id} ---")

    except Exception as e:
        print(f"--- [BG] Processing FAILED: {e} ---")
        traceback.print_exc()
        doc = db.query(CandidateDocument).filter(CandidateDocument.id == document_id).first()
        if doc:
            doc.processing_status = "ERROR"
            doc.processing_error = str(e)
            db.commit()

@router.post('/api/ingest', summary='Ingest File (PDF/TXT) with full pipeline')
async def ingest_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    metadata: str = Form(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Non-blocking Ingestion.
    Returns 202 Accepted immediately.
    """
    try:
        # 0. Parse manual metadata
        try:
            manual_metadata = json.loads(metadata)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid JSON in metadata field")

        print(f"--- Queuing Ingestion for {file.filename} ---")
        
        if not (file.filename.lower().endswith(".pdf") or file.filename.lower().endswith(".docx")):
             raise HTTPException(status_code=400, detail="Only PDF or DOCX supported.")

        # Read file bytes into memory (for passing to bg task)
        # WARNING: Large files might consume memory. For SaaS, stream to S3/Disk is better.
        # For now (MVP/Tool), memory is fine.
        file_bytes = await file.read()

        # 1. Create DB Record (PENDING)
        new_doc = CandidateDocument(
            tenant_id=user.tenant_id,
            filename=file.filename,
            file_type=file.filename.split('.')[-1],
            full_text="", # Will be populated in BG
            candidate_name="Processing...",
            email="",
            processing_status="PENDING"
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)

        # 2. Dispatch Background Task
        # We need a new Session for the background task? 
        # Ideally yes, or we pass the ID and open a new session inside. 
        # The 'db' dependency session closes after request.
        # So we MUST open a new session inside `process_file_background`.
        # OR better: Refactor `process_file_background` to take a session factory or handle it.
        # Since I can't easily pass the factory, I'll instantiate SessionLocal inside the BG function.
        
        # We'll wrap the logic to handle session creation
        background_tasks.add_task(
            run_bg_process, 
            file.filename, 
            file.content_type, 
            file_bytes, 
            manual_metadata, 
            user.tenant_id, 
            new_doc.id
        )

        return {
            "status": "queued",
            "document_id": new_doc.id,
            "message": "File accepted for processing. Check status later."
        }

    except Exception as e:
        print(f"Error queuing ingestion: {e}")
        raise HTTPException(status_code=500, detail=str(e))

from db.database import SessionLocal

async def run_bg_process(filename, content_type, file_bytes, metadata, tenant_id, doc_id):
    # Wrapper to manage DB session lifecycle for BG task
    db = SessionLocal()
    try:
        await process_file_background(filename, content_type, file_bytes, metadata, tenant_id, doc_id, db)
    finally:
        db.close()
